from docx import Document
from docx.shared import RGBColor
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from tkinter import messagebox

class WordDocCreator:

    def __init__(self, document_name: str, checklist_name: str):
        self.document = Document()
        self.document_name = document_name

        heading = self.document.add_heading(f"{checklist_name} Checklist", 0)
        run = heading.runs[0]
        run.font.color.rgb = RGBColor(0, 128, 0)

    def __enter__(self):
        return self

    def __exit__(self, exc_type, exc_val, exc_tb):
        try:
            self.document.save(self.document_name)
            messagebox.showinfo("Success", f"Successfully saved document to {self.document_name}")
        except PermissionError as e:
            messagebox.showerror("Error occurred", f'{e.filename} could not be saved. {e.strerror}')

    def write_output(self, document_contents: list[dict[str, str]]) -> None:
        """Write the content in document_contents to a Word document."""

        table = self.document.add_table(rows=1, cols=2)
        table.style = "Table Grid"

        header_fill_color = "008000"

        header_texts = ("Step Description", "Step Results")
        hdr_cells = table.rows[0].cells
        for i, text in enumerate(header_texts):
            p = hdr_cells[i].paragraphs[0]
            p.clear()
            run = p.add_run(text)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)

            shading_elm = parse_xml(
                r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), header_fill_color)
            )

            hdr_cells[i]._tc.get_or_add_tcPr().append(shading_elm) # No great public API way to make the color green.

        for item in document_contents:
            for key, value in item.items():
                if not key and value:
                    continue

                row_cells = table.add_row().cells
                p_left = row_cells[0].paragraphs[0]
                p_left.clear()
                run_left = p_left.add_run(key or "")
                run_left.font.bold = True

                p_right = row_cells[1].paragraphs[0]
                p_right.clear()
                p_right.add_run(value or "")
